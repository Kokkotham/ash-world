"""
Convert all .docx files in a source directory to structured JSON.

Reads every .docx file, extracts paragraphs (skipping blank lines) and tables
(header row + data rows), then writes each result as a JSON file to a target
output directory.

Usage:
    python docx2json.py [source_dir] [output_dir]

Defaults:
    source_dir = E:/Desktop/跑团文件/规则/EW/
    output_dir = data/ew_raw/
"""

import json
import os
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx is required. Run: pip install python-docx>=1.1.0")
    sys.exit(1)


def extract_paragraphs(doc: Document) -> list[str]:
    """Extract all non-empty paragraph texts from a docx Document."""
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text: str = para.text.strip()
        if text:
            paragraphs.append(text)
    return paragraphs


def extract_tables(doc: Document) -> list[dict]:
    """Extract all tables from a docx Document.

    Each table is converted to:
        {
            "headers": ["col0_header", "col1_header", ...],
            "rows": [
                {"col0": "value", "col1": "value", ...},
                ...
            ]
        }

    If the first row looks like a header (all cells bold or first row only),
    it is treated as the header row. Otherwise, headers are generated as
    "col0", "col1", etc.
    """
    tables: list[dict] = []
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cells: list[str] = [cell.text.strip() for cell in row.cells]
            # Skip fully empty rows
            if any(cells):
                rows.append(cells)

        if not rows:
            continue

        # Determine if first row is a header by checking for bold styling
        first_row = table.rows[0]
        first_row_has_bold: bool = False
        for cell in first_row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.bold:
                        first_row_has_bold = True
                        break

        if first_row_has_bold and len(rows) >= 2:
            headers: list[str] = rows[0]
            data_rows: list[list[str]] = rows[1:]
        else:
            # Generate positional headers
            max_cols: int = max(len(r) for r in rows)
            headers = [f"col{i}" for i in range(max_cols)]
            data_rows = rows

        # Normalise row lengths to match header count
        obj_rows: list[dict] = []
        for row_cells in data_rows:
            row_obj: dict = {}
            for i, header in enumerate(headers):
                row_obj[header] = row_cells[i] if i < len(row_cells) else ""
            obj_rows.append(row_obj)

        tables.append({
            "headers": headers,
            "rows": obj_rows,
        })

    return tables


def convert_docx(file_path: Path) -> dict:
    """Convert a single .docx file to a structured dict."""
    doc: Document = Document(str(file_path))
    return {
        "source": file_path.name,
        "paragraphs": extract_paragraphs(doc),
        "tables": extract_tables(doc),
    }


def main() -> None:
    """Main entry point: scan source dir for .docx, convert each to JSON."""
    source_dir: Path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(
        "E:/Desktop/跑团文件/规则/EW/"
    )
    output_dir: Path = Path(sys.argv[2]) if len(sys.argv) > 2 else Path(
        "data/ew_raw/"
    )

    if not source_dir.exists():
        print(f"ERROR: Source directory does not exist: {source_dir}")
        sys.exit(1)

    output_dir.mkdir(parents=True, exist_ok=True)

    docx_files: list[Path] = sorted(source_dir.glob("*.docx"))
    if not docx_files:
        print(f"WARNING: No .docx files found in {source_dir}")
        return

    for docx_path in docx_files:
        try:
            data: dict = convert_docx(docx_path)
            out_name: str = docx_path.stem + ".json"
            out_path: Path = output_dir / out_name
            with open(out_path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            print(f"OK  {docx_path.name} -> {out_path}")
        except Exception as exc:
            print(f"ERR {docx_path.name}: {exc}")

    print(f"\nDone. {len(docx_files)} file(s) processed.")


if __name__ == "__main__":
    main()
